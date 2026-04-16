#!/usr/bin/env python3
"""Genera documento Word con descripcion operativa, funcional y tecnica de MisReclamos."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def add_title(doc: Document, text: str) -> None:
    title = doc.add_heading(text, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_paragraph(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def main() -> None:
    doc = Document()

    add_title(doc, "MisReclamos - Descripcion operativa, funcional y tecnica")

    add_paragraph(
        doc,
        "Documento tecnico de referencia de la plataforma MisReclamos para operacion diaria, soporte, "
        "evolucion funcional y analisis de arquitectura. Cubre procesos de negocio, flujo de datos, "
        "modelo de seguridad, integraciones y esquema de informacion.",
    )

    add_heading(doc, "1. Proposito y alcance")
    add_paragraph(
        doc,
        "MisReclamos es una plataforma web para gestion integral de reclamos y casos legales. "
        "Centraliza conversaciones de WhatsApp, automatiza tareas con IA, organiza trabajo por estado/prioridad/tipo, "
        "y permite seguimiento operativo por dashboard y trazabilidad completa del expediente digital.",
    )
    add_bullets(
        doc,
        [
            "Canal principal: WhatsApp (integracion BuilderBot Cloud).",
            "Backoffice legal: administracion de tickets, clientes, abogados, tipos de caso y configuracion IA.",
            "Base de datos transaccional: PostgreSQL con Prisma ORM.",
            "Adjuntos y multimedia: almacenamiento persistente en Vercel Blob.",
            "Capas de IA: resumen de conversacion, clasificacion legal, inferencia de prioridad y transcripcion de audios.",
        ],
    )

    add_heading(doc, "2. Arquitectura tecnica")
    add_paragraph(doc, "La aplicacion esta implementada en Next.js (App Router) con TypeScript.")
    add_bullets(
        doc,
        [
            "Capa UI: rutas server/client en Next.js para panel operativo.",
            "Capa API: endpoints REST en /api/* para operaciones de negocio y administracion.",
            "Persistencia: Prisma Client contra PostgreSQL.",
            "Sesion y autenticacion: iron-session en cookie httpOnly (misreclamos_session).",
            "Integraciones externas: BuilderBot, OpenAI API y Vercel Blob.",
        ],
    )

    add_heading(doc, "3. Modelo de usuarios, roles y control de acceso")
    add_paragraph(
        doc,
        "Existen dos roles de operacion: ADMIN y SUPPORT. El alcance de lectura y escritura sobre tickets "
        "se define por rol y asignacion.",
    )
    add_bullets(
        doc,
        [
            "ADMIN: visibilidad total, reasignacion de casos, gestion de clientes/agentes/casos/configuracion/usuarios.",
            "SUPPORT: visibilidad limitada a tickets asignados al propio usuario.",
            "Las rutas de gestion administrativa validan rol ADMIN en backend.",
            "El acceso a tickets individuales valida autorizacion por usuario/rol antes de exponer datos.",
        ],
    )

    add_heading(doc, "4. Inicio de sesion y ciclo de credenciales")
    add_bullets(
        doc,
        [
            "Login principal: username + password (hash bcrypt en AgentUser.passwordHash).",
            "Normalizacion de usuario: trim + minusculas.",
            "Primer acceso: endpoint dedicado para bootstrap de primer administrador cuando no existen contrasenas.",
            "Modo legacy opcional: usuario admin + APP_PASSWORD (solo respaldo de migracion).",
            "Sesion emitida en cookie segura (secure en produccion, sameSite=lax).",
        ],
    )

    add_heading(doc, "5. Entidades de datos principales")
    add_paragraph(doc, "El modelo de dominio se apoya en las siguientes entidades:")
    add_bullets(
        doc,
        [
            "Customer: telefono unico, datos personales y estado botPausedAt para moderacion del bot.",
            "AgentUser: identidad del operador, rol, credenciales y datos profesionales.",
            "Ticket: caso/reclamo con codigo unico, estado, prioridad, categoria, canal, tipo legal y metadatos procesales.",
            "TicketMessage: historial conversacional con direccion, origen, texto, adjuntos, payload e id externo.",
            "TicketEvent: bitacora de eventos operativos (escalado, cambio de estado, auto reply, etc.).",
            "CaseType: catalogo dinamico de tipos de caso (slug, label, legalType, color, orden).",
            "Legado: almacenamiento de registros historicos importados desde Excel.",
        ],
    )

    add_heading(doc, "6. Flujo operativo principal (WhatsApp -> Ticket)")
    add_paragraph(
        doc,
        "El webhook /api/whatsapp/inbound recibe eventos de BuilderBot y orquesta el procesamiento completo.",
    )
    add_bullets(
        doc,
        [
            "Parseo y validacion del payload entrante (message.incoming / message.outgoing).",
            "Upsert de cliente por telefono.",
            "Deteccion de ticket activo por ventana temporal (48h) o creacion de nuevo ticket.",
            "Persistencia de mensaje con adjuntos normalizados y payload crudo para auditoria.",
            "Heuristicas de prioridad/categoria y decision de escalado.",
            "Actualizacion de estado del ticket y eventos de negocio.",
            "Auto-respuesta condicional (si no hay pausa de bot y se cumplen reglas de escalado).",
            "Envio o recordatorio de codigo de reclamo al cierre conversacional.",
            "Post-procesamiento IA: resumen y clasificacion legal cuando aplica.",
        ],
    )

    add_heading(doc, "7. Gestion de adjuntos y multimedia")
    add_bullets(
        doc,
        [
            "Archivos entrantes/salientes se descargan y suben a Vercel Blob para URL persistente.",
            "Se exige URL absoluta para evitar referencias invalidas o rotas.",
            "Soporte de imagen, video, audio, pdf y documentos ofimaticos.",
            "En mensajes manuales del backoffice se permite upload directo (multipart/form-data).",
            "En errores de upload, el sistema evita persistir rutas relativas y retorna falla controlada.",
        ],
    )

    add_heading(doc, "8. IA aplicada a la operacion")
    add_bullets(
        doc,
        [
            "Resumen de conversacion: generado con OpenAI (modelo gpt-4o-mini) con fallback sin API key.",
            "Clasificacion de tipo legal: seleccion exacta entre categorias permitidas.",
            "Inferencia de prioridad: LOW/NORMAL/HIGH/URGENT segun contenido conversacional.",
            "Transcripcion de notas de voz: Whisper (whisper-1) en idioma espanol.",
            "Filtro de mensajes no resumibles: placeholders de eventos y mensajes vacios/adjuntos sin texto.",
        ],
    )

    add_heading(doc, "9. Bandeja de casos y operacion diaria")
    add_bullets(
        doc,
        [
            "Listado central ordenado por ultima actividad.",
            "Busqueda por codigo, titulo, cliente y contacto.",
            "Filtros por estado, prioridad y tipo de caso.",
            "Contadores operativos (abiertos, en progreso, esperando cliente, urgentes, resueltos, cerrados).",
            "Detalle de caso con timeline completo, composicion de mensajes, adjuntos y acciones rapidas.",
            "Cambio de estado, asignacion de abogado y actualizacion de tipo legal desde UI.",
            "Refresco en vivo para reducir latencia operativa en conversaciones activas.",
        ],
    )

    add_heading(doc, "10. Dashboard y observabilidad funcional")
    add_bullets(
        doc,
        [
            "KPI principales: total de casos, creados hoy, resueltos hoy y tiempo promedio de resolucion.",
            "Distribucion por estado y prioridad.",
            "Tendencia de nuevos casos en ultimos 7 dias.",
            "Alerta de casos urgentes activos con acceso directo al filtro.",
            "Top de abogados por carga activa y top de clientes por volumen.",
            "Desglose por tipo legal con navegacion directa a vista filtrada.",
            "Respeta alcance por rol (ADMIN global, SUPPORT acotado a asignados).",
        ],
    )

    add_heading(doc, "11. Modulos administrativos")
    add_bullets(
        doc,
        [
            "Clientes: alta manual, busqueda, paginacion, importacion Excel, edicion y eliminacion.",
            "Abogados/Agentes: alta, actualizacion, baja y medicion de carga asignada.",
            "Tipos de caso: CRUD del catalogo dinamico usado por menu y filtros.",
            "Usuarios de acceso: asignacion y reseteo de username/password a AgentUser.",
            "Configuracion IA: edicion del prompt del bot en BuilderBot.",
            "Base de conocimiento: carga/eliminacion de archivos de entrenamiento en BuilderBot.",
        ],
    )

    add_heading(doc, "12. Notificaciones y coordinacion interna")
    add_paragraph(
        doc,
        "Al reasignar un ticket a un abogado, la plataforma puede notificar por WhatsApp al profesional asignado "
        "con resumen del caso y datos operativos relevantes. El fallo de esta notificacion no bloquea la operacion core.",
    )

    add_heading(doc, "13. Seguridad, confiabilidad y resiliencia")
    add_bullets(
        doc,
        [
            "Validacion de payloads con Zod en endpoints criticos.",
            "Control de acceso en backend (nunca solo UI).",
            "Hash de contrasenas con bcrypt.",
            "Idempotencia por externalMessageId para minimizar duplicados.",
            "Manejo defensivo de fallos en integraciones externas (respuestas de fallback).",
            "Logs operativos para trazabilidad de eventos y errores.",
        ],
    )

    add_heading(doc, "14. Integraciones externas y dependencias")
    add_bullets(
        doc,
        [
            "BuilderBot Cloud: mensajeria, prompt del agente y archivos de conocimiento.",
            "OpenAI: resumen, clasificacion, prioridad y transcripcion.",
            "Vercel Blob: almacenamiento de adjuntos persistentes.",
            "PostgreSQL + Prisma: almacenamiento transaccional del dominio.",
        ],
    )

    add_heading(doc, "15. Variables de entorno clave")
    add_bullets(
        doc,
        [
            "DATABASE_URL, SESSION_PASSWORD, APP_PASSWORD.",
            "BUILDERBOT_BOT_ID, BUILDERBOT_API_KEY, BUILDERBOT_BASE_URL, BUILDERBOT_API_URL.",
            "BUILDERBOT_BOT_URL y BUILDERBOT_DASHBOARD_TOKEN (escenario self-hosted blacklist).",
            "OPENAI_API_KEY.",
            "BLOB_READ_WRITE_TOKEN.",
            "NEXT_PUBLIC_GESTION_ANALISIS_URL (enlace opcional a app externa).",
        ],
    )

    add_heading(doc, "16. Flujo de despliegue y operacion de base de datos")
    add_bullets(
        doc,
        [
            "El proyecto versiona migraciones Prisma en repositorio.",
            "Railway/Vercel requieren gestionar variables de entorno por ambiente.",
            "Las migraciones pueden aplicarse manualmente via SQL o con deploy de Prisma.",
            "Existe tooling de soporte para chequeo de DB, bootstrap de admin y carga legado.",
        ],
    )

    add_heading(doc, "17. Riesgos operativos y recomendaciones")
    add_bullets(
        doc,
        [
            "Mantener credenciales de integraciones rotadas y auditadas.",
            "Evitar dependencia del modo legacy de login en ambientes productivos.",
            "Monitorear latencia/fallos de BuilderBot y OpenAI para evitar degradacion silenciosa.",
            "Revisar periodicamente reglas heuristicas de escalado para ajustar precision.",
            "Definir backup y retencion de base de datos y archivos adjuntos.",
        ],
    )

    add_heading(doc, "18. Resumen ejecutivo tecnico")
    add_paragraph(
        doc,
        "MisReclamos funciona como una plataforma legal-operativa omnicanal centrada en WhatsApp, "
        "con backoffice robusto para gestion de casos, capa de automatizacion con IA y controles de "
        "acceso por rol. Su arquitectura combina simplicidad de despliegue con capacidad de trazabilidad, "
        "permitiendo operar reclamos en tiempo real con gobierno de datos y soporte administrativo completo.",
    )

    out = (
        Path(__file__).resolve().parent.parent
        / "Descripcion-Operativa-Funcional-Tecnica-MisReclamos.docx"
    )
    doc.save(out)
    print(f"Generado: {out}")


if __name__ == "__main__":
    main()
