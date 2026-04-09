#!/usr/bin/env python3
"""Genera documento Word con descripción funcional del área de contact center (backoffice Mis Reclamos)."""

from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    return p


def add_para(doc: Document, text: str, bold: bool = False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    return p


def main():
    doc = Document()
    title = doc.add_heading(
        "Mis Reclamos — Funcionalidades del área de contacto y atención (backoffice)",
        0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_para(
        doc,
        "Este documento describe las capacidades de la plataforma orientadas al trabajo diario de "
        "atención al cliente y gestión de reclamos. En el código no existe un módulo con el nombre "
        "«Contact Center»; el equivalente operativo es el panel web de Mis Reclamos (casos/reclamos, "
        "conversaciones por WhatsApp, agentes y clientes), con roles de Administrador y Abogado.",
    )

    add_heading(doc, "1. Alcance y concepto", 1)
    add_para(
        doc,
        "El «centro de contacto» en Mis Reclamos agrupa: ingreso y seguimiento de reclamos como tickets, "
        "comunicación multicanal centrada en WhatsApp (integración BuilderBot), asignación a abogados, "
        "moderación del bot automático, resúmenes con IA y paneles de supervisión. Quedan fuera de este "
        "alcance descriptivo las apps externas opcionales (p. ej. «Gestión y análisis» si está configurada) "
        "y la carga histórica «Legado», salvo mención breve al final.",
    )

    add_heading(doc, "2. Roles y permisos", 1)
    add_para(
        doc,
        "Administrador (ADMIN): ve todos los casos, puede asignar y reasignar abogados, gestionar clientes, "
        "abogados, tipos de caso, configuración del agente de IA, base de conocimiento y usuarios de acceso.",
    )
    add_para(
        doc,
        "Abogado / Soporte (SUPPORT): solo accede a los tickets que tiene asignados. No puede cambiar la "
        "asignación del caso a otro colega (eso lo reserva el administrador).",
    )

    add_heading(doc, "3. Acceso y sesión", 1)
    add_para(
        doc,
        "Los usuarios del equipo inician sesión en el backoffice. La sesión identifica nombre y rol. "
        "Existe flujo de primer acceso / definición de credenciales según la política implementada en el proyecto.",
    )

    add_heading(doc, "4. Interfaz principal (layout)", 1)
    add_para(
        doc,
        "Barra lateral con navegación: Dashboard, listado de todos los casos, filtros por estado y por tipo "
        "de caso (tipos dinámicos desde base de datos), y sección de gestión para administradores. "
        "Diseño adaptable con menú hamburguesa en móvil. Cierre de sesión desde el pie del menú.",
    )

    add_heading(doc, "5. Dashboard", 1)
    bullets = [
        "Indicadores: total de casos en alcance del usuario, casos creados hoy, resueltos hoy, tiempo promedio "
        "de resolución (en horas).",
        "Distribución de casos por estado (abierto, en progreso, esperando cliente, resuelto, cerrado) y por prioridad.",
        "Gráfico de tendencia de los últimos 7 días (altas diarias).",
        "Ranking de abogados con más casos activos (no cerrados); para administrador muestra varios, para abogado "
        "muestra su propia posición.",
        "Clientes con más reclamos (top), respetando el alcance del rol.",
        "Desglose de casos por tipo legal (legalType) con enlaces al filtro por tipo.",
        "Alerta destacada si hay casos urgentes aún activos, con acceso directo al listado filtrado de urgentes.",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    add_heading(doc, "6. Bandeja de casos (tickets)", 1)
    add_para(
        doc,
        "Vista central de todos los reclamos visibles para el usuario (hasta un límite razonable de registros "
        "ordenados por última actividad).",
    )
    sub = [
        "Búsqueda por código de reclamo, título, nombre del cliente o nombre de contacto.",
        "Filtros por estado (incluyendo filtro de prioridad urgente) y por tipo de caso (legalType).",
        "Ordenación por último mensaje, fecha de creación o código.",
        "Contadores rápidos: abiertos, en progreso, esperando cliente, urgentes, resueltos y cerrados.",
        "Notas internas por caso: ventana modal para editar y guardar notas del abogado (visibles solo al equipo).",
        "Acceso al detalle de cada ticket.",
    ]
    for s in sub:
        doc.add_paragraph(s, style="List Bullet")

    add_heading(doc, "7. Vistas por estado y por tipo", 1)
    add_para(
        doc,
        "Rutas dedicadas en el menú lateral para filtrar casos: Abiertos, En progreso, Esperando cliente, "
        "Resueltos, Cerrados y Urgentes. Además, cada tipo de caso configurado en el sistema tiene su URL "
        "(/tickets/tipo/[slug]) para trabajar solo esa matrícula de asuntos.",
    )

    add_heading(doc, "8. Detalle de un reclamo (ticket)", 1)
    add_para(doc, "Cabecera:", bold=True)
    add_para(
        doc,
        "Código único del ticket, persona que contacta, tipo de caso, teléfono formateado, acciones de cambio "
        "de estado, asignación de abogado (según permisos), botón para ir al hilo de conversación y responder.",
    )
    add_para(doc, "Hilo de conversación:", bold=True)
    add_para(
        doc,
        "Muestra mensajes en orden cronológico con distinción visual entre Cliente, Bot y Agente humano. "
        "Soporta adjuntos (imágenes, documentos, etc.) con visualización/descarga según lo almacenado.",
    )
    add_para(doc, "Compositor de mensajes:", bold=True)
    add_para(
        doc,
        "El agente puede enviar: respuesta al cliente por WhatsApp (saliente) o nota interna (no enviada al "
        "cliente). Permite adjuntar archivos (límite de tamaño configurable en la UI) con tipos admitidos "
        "(PDF, Office, imágenes, audio, vídeo, etc.). Tras enviar, la vista se actualiza.",
    )
    add_para(doc, "Pausar agente (bot) por cliente:", bold=True)
    add_para(
        doc,
        "Desde el detalle se puede pausar o reactivar el bot para ese cliente: cuando está pausado, el sistema "
        "no envía auto-respuestas automáticas del bot y el equipo responde manualmente; la acción sincroniza "
        "lista negra en BuilderBot (cloud y self-hosted según configuración).",
    )
    add_para(doc, "Panel lateral — datos y acciones:", bold=True)
    side = [
        "Datos del cliente con enlace directo a WhatsApp (wa.me) si hay teléfono.",
        "Selector de tipo de caso legal (legalType) alineado con el catálogo de tipos de caso.",
        "Asignación de abogado (duplicada en cabecera y lateral; reasignación solo administrador).",
        "Resumen inteligente: texto generado por IA a partir de la conversación, con botón para generar o regenerar. "
        "El webhook de entrada también puede actualizar este resumen automáticamente cuando hay nueva actividad.",
        "Bloque de acciones rápidas orientativo (notas, estado, respuesta).",
    ]
    for s in side:
        doc.add_paragraph(s, style="List Bullet")

    add_para(doc, "Actualización en tiempo casi real:", bold=True)
    add_para(
        doc,
        "El detalle incluye mecanismo de refresco para que los mensajes nuevos se perciban sin depender solo "
        "de recargar la página manualmente.",
    )

    add_heading(doc, "9. Estados, prioridades y modelo de datos del caso", 1)
    add_para(
        doc,
        "Estados: Abierto, En progreso, Esperando cliente, Resuelto, Cerrado. Prioridades: Baja, Normal, Alta, Urgente. "
        "Categorías legales amplias (laboral, civil, etc.) y canal (WhatsApp, email, web) existen en el modelo para "
        "clasificación. El ticket incluye además campos procesales opcionales (expediente, juzgado, montos, fechas "
        "importantes, resolución, etc.) según evolución del producto.",
    )

    add_heading(doc, "10. Canal WhatsApp e integración BuilderBot", 1)
    add_para(
        doc,
        "Webhook de mensajes entrantes y salientes: registra en el ticket los mensajes del cliente y del flujo "
        "del bot; procesa adjuntos subiéndolos a almacenamiento persistente (blob); maneja eventos de multimedia "
        "y notas de voz, con transcripción automática cuando aplica.",
    )
    add_para(
        doc,
        "Lógica de negocio: creación o vinculación de cliente por teléfono, generación de código de reclamo, "
        "respuestas automáticas en escenarios definidos (p. ej. escalamiento cuando el usuario pide hablar con "
        "un agente), envío del código de reclamo en despedidas, actualización de resumen con IA y clasificación "
        "inicial del tipo de caso si aún no estaba definido.",
    )
    add_para(
        doc,
        "Mensajes salientes desde el panel se envían al cliente vía la integración de WhatsApp; los duplicados "
        "por doble registro (bot + humano) se gestionan en backend para evitar repeticiones.",
    )

    add_heading(doc, "11. Notificación al abogado al asignar un caso", 1)
    add_para(
        doc,
        "Cuando un administrador asigna un ticket a un abogado, el sistema puede enviar un mensaje de WhatsApp "
        "al teléfono profesional del abogado con resumen del caso, datos clave y enlace al panel (según URL configurada).",
    )

    add_heading(doc, "12. Gestión de clientes (solo administrador)", 1)
    add_para(
        doc,
        "Alta manual de clientes, importación desde Excel, listado con paginación/cantidad reciente, datos de "
        "contacto y moderación del bot (pausa) coherente con la ficha en reclamos.",
    )

    add_heading(doc, "13. Gestión de abogados / agentes de soporte (solo administrador)", 1)
    add_para(
        doc,
        "Alta de usuarios del equipo con rol de soporte, datos profesionales (matrícula, especialización), teléfono "
        "para notificaciones, listado con conteo de casos asignados y métricas simples de actividad.",
    )

    add_heading(doc, "14. Tipos de caso (solo administrador)", 1)
    add_para(
        doc,
        "Administración del catálogo de tipos: etiqueta visible, slug para URLs, valor legalType asociado al ticket, "
        "color en la interfaz y orden. Esto alimenta el menú lateral «Por tipo de caso» y los filtros del dashboard.",
    )

    add_heading(doc, "15. Configuración del agente de IA y base de conocimiento (solo administrador)", 1)
    add_para(
        doc,
        "Prompt editable del asistente conectado a BuilderBot (instrucciones de tono y comportamiento). "
        "Gestión de archivos de conocimiento para el sistema según la implementación de la pantalla de configuración.",
    )

    add_heading(doc, "16. Usuarios de acceso (solo administrador)", 1)
    add_para(
        doc,
        "Administración de credenciales y usuarios con acceso al backoffice (según pantalla de usuarios en configuración).",
    )

    add_heading(doc, "17. Funciones colaterales no centrales en atención", 1)
    add_para(
        doc,
        "Legado: importación de filas históricas desde Excel para referencia; no es el flujo principal de contact center. "
        "Enlace externo «Gestión y análisis»: solo aparece si la variable de entorno pública está definida; es una "
        "aplicación aparte.",
    )

    add_heading(doc, "18. Resumen", 1)
    add_para(
        doc,
        "El núcleo del contact center en Mis Reclamos es la gestión unificada de conversaciones y reclamos legales "
        "desde el backoffice, con fuerte integración WhatsApp/BuilderBot, trabajo por colas (estado, urgencia, tipo), "
        "asignación a abogados, apoyo de IA en resúmenes y clasificación, y paneles de seguimiento para supervisión "
        "y mejora continua.",
    )

    out = Path(__file__).resolve().parent.parent / "Funcionalidades-Contact-Center-MisReclamos.docx"
    doc.save(out)
    print(f"Generado: {out}")


if __name__ == "__main__":
    main()
